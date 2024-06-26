import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from transformers import AutoTokenizer , AutoModelForSequenceClassification,pipeline
import io,base64,re
import streamlit_pills as stp
from typing import List
from stqdm import stqdm
import time


LABELS = ['Value equation',
 'Credentialing / Quality Assurance Infrastructure',
 'Finanicial Impact',
 'Health System Characteristics',
 'Clinical utility & efficiency-Provider perspective',
 'Workflow related problems',
 'Provider Characteristics',
 'Training',
 'Patient/Physican interaction in LUS',
 'Imaging modalities in general']

TOP_LEVELS = [
    "Multi-level organizations Characteristics-creating an environment (infrastructure) for encouraging spread",
    "Multi-level organizations Perspectives/Values -sharing best practices; observing results and adjusting processes accordingly",
    "Implementation and Sustainability infrastructure- facilitating use of the intervention; -ensuring adaptability of protocols that fit the multilevel context"
]


multi_level_org_char: List[str] = ['Provider characteristics', "Health System Characteristics"]
multi_level_org_perspect: List[str] = [ "Imaging modalities in general",
                            'Value equation',
                            "Clinical utility & efficiency-Provider perspective",
                            "Patient/Physican interaction in LUS",
                            'Workflow related problems']

impl_sust_infra: List[str] = [
    "Training",
    'Credentialing / Quality Assurance Infrastructure',
    "Finanicial Impact",
]




def get_binary_file_downloader_html(bin_file, file_label='File'):
    bin_file.seek(0)
    encoded_file = base64.b64encode(bin_file.read()).decode()
    bin_file.close()
    href = f'<a href="data:application/octet-stream;base64,{encoded_file}" download="processed_document.docx">{file_label}</a>'
    return href



def classify_new_text(text:str , model_path):
    classifier = model_path
    output = classifier(text)

    if output[0]['label'] in multi_level_org_char:
        output[0]['top level'] = TOP_LEVELS[0]
    elif output[0]['label'] in multi_level_org_perspect:
        output[0]['top level'] = TOP_LEVELS[1]
    elif output[0]['label'] in impl_sust_infra:
        output[0]['top level'] = TOP_LEVELS[2]

    return output[0]


COLOR_LIST = [WD_COLOR_INDEX.DARK_BLUE,
              WD_COLOR_INDEX.RED,
              WD_COLOR_INDEX.DARK_YELLOW,
              WD_COLOR_INDEX.TURQUOISE,
              WD_COLOR_INDEX.VIOLET,
              WD_COLOR_INDEX.YELLOW,
              WD_COLOR_INDEX.PINK,
              WD_COLOR_INDEX.DARK_RED,
              WD_COLOR_INDEX.GRAY_50,
              WD_COLOR_INDEX.GREEN,
              ]

low_color_dict = dict(zip(LABELS, COLOR_LIST))
high_color_dict = dict(zip(TOP_LEVELS, [RGBColor(255, 0, 0),RGBColor(0, 255, 0),RGBColor(0, 0, 255)])) # red, green and blue for the top 3 levels



def apply_low_highlight(paragraph, label):
    color_index = low_color_dict.get(label, WD_COLOR_INDEX.AUTO)
    run.font.highlight_color = color_index

def apply_high_highlight(paragraph,label):
    color_index = high_color_dict.get(label, RGBColor(0,0,0))
    run.font.color.rgb = color_index





st.title("Healthcare Document Classification and Highlighting")


MODEL_CHOICE = stp.pills('Choose the model to use', ['Distilbert', 'Electra'], index=None)

if MODEL_CHOICE=='Distilbert':
    with st.spinner("Loading bert weights..."):
        time.sleep(1)
        bert_model = pipeline('text-classification', 'eskayML/interview_classifier', tokenizer = AutoTokenizer.from_pretrained('eskayML/interview_classifier'))

    
    
elif MODEL_CHOICE=='Electra':
    with st.spinner("Loading electra weights..."):
        time.sleep(1)
        electra_model = pipeline('text-classification', 'eskayML/interview_electra', tokenizer = AutoTokenizer.from_pretrained('eskayML/interview_electra'))


def classify(paragraph, MODEL_CHOICE):
    output = classify_new_text(paragraph, bert_model) if MODEL_CHOICE == 'Distilbert' else classify_new_text(paragraph, electra_model)

    thres_zip = dict(zip(['Distilbert', 'Electra'], [0.5,0.12]))

    if output['score'] > thres_zip[MODEL_CHOICE]:
        return output
    

# Load the document
file_upload = st.file_uploader("Upload a Word document (.docx)", type=["docx"])

if file_upload:

    if not MODEL_CHOICE:
        st.error('No model selected!!')
        exit()

    uploaded_file = file_upload.name
    doc = Document(file_upload)

    high_label_counts = {label: 0 for label in high_color_dict.keys()}
    low_label_counts = {label: 0 for label in low_color_dict.keys()}

    # Process paragraphs and apply highlights to chosen sentences
    for paragraph in stqdm(doc.paragraphs, desc="Classifying Sentences..."):
        
        sentences = re.split(r'\.\s*', paragraph.text)
        for sentence in sentences:
            words = sentence.split()
            if len(words) > 10:  # Only consider sentences with more than 10 words
                chosen_sentence = ' '.join(words)
                
                label = classify(chosen_sentence, MODEL_CHOICE)
                if label:
                    high_label = label['top level']
                    low_label = label['label']
                    high_label_counts[high_label] += 1
                    low_label_counts[low_label] += 1
                    run = paragraph.add_run(chosen_sentence)  # Create a new run for the chosen sentence
                    apply_high_highlight(run, high_label)
                    apply_low_highlight(run, low_label)

            
    legend_paragraph = doc.add_paragraph("LEGEND: TOP LEVEL COLOR IDENTIFICATION")
    for label, color_index in high_color_dict.items():
        run = legend_paragraph.add_run()
        run.text = f"\n{label}:"
        run.font.color.rgb = color_index

    legend_paragraph = doc.add_paragraph("\nLEGEND: SUB LEVEL COLOR IDENTIFICATION")
    for label, color_index in low_color_dict.items():
        run = legend_paragraph.add_run()
        run.text = f"\n{label}:"
        run.font.highlight_color  = color_index


    # Add summary at the end of the document
    summary_paragraph = doc.add_paragraph("\n\nSUMMARY:\n")

    summary_paragraph.add_run("\nHigh-level label counts:\n").bold = True
    for label, count in high_label_counts.items():
        summary_paragraph.add_run(f"{label}: {count}\n")



    summary_paragraph.add_run("\nLow-level label counts:\n").bold = True
    for label, count in low_label_counts.items():
        summary_paragraph.add_run(f"{label}: {count}\n")


    max_high_label = max(high_label_counts, key=high_label_counts.get)
    max_low_label = max(low_label_counts, key=low_label_counts.get)

    # Calculate the percentage of occurrence for the high-level label
    total_high_labels = sum(high_label_counts.values())
    try:
        percentage_high_label = (high_label_counts[max_high_label] / total_high_labels) * 100
    except ZeroDivisionError:
        percentage_high_label = 0

    # Calculate the percentage of occurrence for the low-level label
    total_low_labels = sum(low_label_counts.values())
    try:
        percentage_low_label = (low_label_counts[max_low_label] / total_low_labels) * 100
    except ZeroDivisionError:
        percentage_low_label = 0


    # Add summary about the label that occurred the most
    summary_paragraph.add_run("\nMost Occurred High-level Label:\n").bold = True
    summary_paragraph.add_run(f"{max_high_label}: {percentage_high_label:.2f}%\n")

    summary_paragraph.add_run("\nMost Occurred Low-level Label:\n").bold = True
    summary_paragraph.add_run(f"{max_low_label}: {percentage_low_label:.2f}%\n")
    # Save the document with a new name
    output_file = io.BytesIO()
    doc.save(output_file)
    output_file.seek(0)


    # Provide a download link for the output file
    st.markdown(get_binary_file_downloader_html(output_file, file_label="Download Processed Document"), unsafe_allow_html=True)

