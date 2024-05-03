import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from transformers import AutoTokenizer , AutoModelForSequenceClassification,pipeline
import io,base64


LABELS = ['Value equation: quality/cost/efficiency/patient satisfaction (Hosp leaders-background for implementation)',
 'credientialing /quality assurance infrastructure',
 'Finanicial Impact',
 'Health System characteristics (academics; bed capacity also)',
 'Clinical utility & efficiency-Provider perspective',
 'Workflow (with subcoding [bolded]: access equipment, order vs encounter-based; uploading & saving images; type of equipment.  End user need to acquire and use',
 'Provider characteristcs',
 'training',
 'Patient/Physican interaction in LUS',
 'Imaging modalities in general',
 'Clinical utility &  efficiency-Provider perspective',
 'Value equation: quality/cost/efficiency/patient satisfaction (Hosp leaders) High-value care (providers)'
 ]


TOP_LEVELS = [
    "Multi-level organizations Characteristics-creating an environment (infrastructure) for encouraging spread ",
    "Multi-level organizations Perspectives/Values -sharing best practices; observing results and adjusting processes accordingly",
    "Implementation and Sustainability infrastructure- facilitating use of the intervention; -ensuring adaptability of protocols that fit the multilevel context"
]


model_path  = 'eskayML/interview_classifier'

tokenizer  =  AutoTokenizer.from_pretrained('distilbert-base-uncased')

multi_level_org_char = ['Provider characteristcs', "Health System characteristics (academics; bed capacity also)"]
multi_level_org_perspect = [ "Imaging modalities in general",
                            "Value equation: quality/cost/efficiency/patient satisfaction (Hosp leaders) High-value care (providers)",
                            "Clinical utility &  efficiency-Provider perspective",
                            "Patient/Physican interaction in LUS",
                            "Workflow (with subcoding [bolded]: access equipment, order vs encounter-based; uploading & saving images; type of equipment.  End user need to acquire and use"]

impl_sust_infra = [
    "training",
    "credientialing /quality assurance infrastructure",
    "Finanicial Impact",
]



def classify_new_text(text):
  classifier = pipeline('text-classification', model_path, tokenizer = tokenizer)
  output = classifier(text)

  if output[0]['label'] in multi_level_org_char:
    output[0]['top level'] = TOP_LEVELS[0]
  elif output[0]['label'] in multi_level_org_perspect:
    output[0]['top level'] = TOP_LEVELS[1]
  elif output[0]['label'] in impl_sust_infra:
    output[0]['top level'] = TOP_LEVELS[2]

  return output



COLOR_LIST = [WD_COLOR_INDEX.BRIGHT_GREEN,
              WD_COLOR_INDEX.RED,
              WD_COLOR_INDEX.TEAL,
              WD_COLOR_INDEX.TURQUOISE,
              WD_COLOR_INDEX.VIOLET,
              WD_COLOR_INDEX.YELLOW,
              WD_COLOR_INDEX.PINK,
              WD_COLOR_INDEX.DARK_RED,
              WD_COLOR_INDEX.GRAY_50,
              WD_COLOR_INDEX.BLUE,
              WD_COLOR_INDEX.DARK_YELLOW,
              WD_COLOR_INDEX.DARK_BLUE,
              ]

low_color_dict = dict(zip(LABELS, COLOR_LIST))
high_color_dict = dict(zip(TOP_LEVELS, [RGBColor(255, 0, 0),RGBColor(0, 255, 0),RGBColor(0, 0, 255)])) # red, green and blue for the top 3 levels


def get_binary_file_downloader_html(bin_file, file_label='File'):
    bin_file.seek(0)
    encoded_file = base64.b64encode(bin_file.read()).decode()
    bin_file.close()
    href = f'<a href="data:application/octet-stream;base64,{encoded_file}" download="processed_document.docx">{file_label}</a>'
    return href

def classify(paragraph):
    output = classify_new_text(paragraph)
    return output[0] if output[0]['score'] > 0.5 else None

def apply_low_highlight(paragraph, label):
    color_index = low_color_dict.get(label, WD_COLOR_INDEX.AUTO)
    for run in paragraph.runs:
        run.font.highlight_color = color_index

def apply_high_highlight(paragraph,label):
    color_index = high_color_dict.get(label, RGBColor(0,0,0))
    for run in paragraph.runs:
        run.font.color.rgb = color_index


def main():
    st.title("Interview Text Classification and Highlighting")

    # Load the document
    file_upload = st.file_uploader("Upload a Word document (.docx)", type=["docx"])

    if file_upload:
        uploaded_file = file_upload.name
        doc = Document(file_upload)
        # Add legends at the top of the document
        legend_paragraph = doc.add_paragraph("LEGEND::=> TOP LEVEL COLOR IDENTIFICATION")
        for label, color_index in high_color_dict.items():
            run = legend_paragraph.add_run()
            run.text = f"\n{label}:"
            run.font.color.rgb = color_index

        legend_paragraph = doc.add_paragraph("LEGEND::=> SUB LEVEL COLOR IDENTIFICATION")
        for label, color_index in low_color_dict.items():
            run = legend_paragraph.add_run()
            run.text = f"\n{label}:"
            run.font.highlight_color  = color_index


        # ... (rest of the code remains the same)
        high_label_counts = {label: 0 for label in high_color_dict.keys()}
        low_label_counts = {label: 0 for label in low_color_dict.keys()}


        # Process paragraphs and apply highlights
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith('Interviewee:'):
                text = paragraph.text[len('Interviewee:'):].strip()

                label = classify(text)
                if label:
                    high_label = label['top level']
                    low_label = label['label']
                    high_label_counts[high_label] += 1
                    low_label_counts[low_label] += 1
                    apply_high_highlight(paragraph, high_label)
                    apply_low_highlight(paragraph, low_label)

        # Add summary at the end of the document
        summary_paragraph = doc.add_paragraph("SUMMARY:\n")

        summary_paragraph.add_run("High-level label counts:\n").bold = True
        for label, count in high_label_counts.items():
            summary_paragraph.add_run(f"{label}: {count}\n")

        summary_paragraph.add_run("\nLow-level label counts:\n").bold = True
        for label, count in low_label_counts.items():
            summary_paragraph.add_run(f"{label}: {count}\n")

        max_high_label = max(high_label_counts, key=high_label_counts.get)
        max_low_label = max(low_label_counts, key=low_label_counts.get)

        # Calculate the percentage of occurrence for the high-level label
        total_high_labels = sum(high_label_counts.values())
        percentage_high_label = (high_label_counts[max_high_label] / total_high_labels) * 100

        # Calculate the percentage of occurrence for the low-level label
        total_low_labels = sum(low_label_counts.values())
        percentage_low_label = (low_label_counts[max_low_label] / total_low_labels) * 100

        # Add summary about the label that occurred the most
        summary_paragraph.add_run("\nMost Occurred High-level Label:\n").bold = True
        summary_paragraph.add_run(f"{max_high_label}: {percentage_high_label:.2f}%\n")

        summary_paragraph.add_run("\nMost Occurred Low-level Label:\n").bold = True
        summary_paragraph.add_run(f"{max_low_label}: {percentage_low_label:.2f}%\n")

        # Save the document with a new name
        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)

        # Provide a download link for the output file
        st.markdown(get_binary_file_downloader_html(output_file, file_label="Download Processed Document"), unsafe_allow_html=True)


if __name__ == "__main__":
    main()
